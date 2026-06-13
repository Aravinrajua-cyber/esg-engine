from __future__ import annotations

import pickle
from pathlib import Path

import pandas as pd
from docx import Document

from src.report.build_report import CLAUDE_MARKER, _load_phase4_results, _phase4_tables, build_final_report


def test_phase4_pickle_tables_are_loaded(tmp_path: Path):
    phase4_path = tmp_path / "phase4_results.pkl"
    payload = {
        "ic_results": pd.DataFrame(
            [{"variable": "sentiment_dynamics", "horizon": 3, "ic_mean": 0.12, "t_nw": 2.1, "fdr_survived": True}]
        ),
        "fama_macbeth_coefficients": pd.DataFrame(
            [{"variable": "sentiment_dynamics", "coef": 0.04, "t_nw": 2.4, "p": 0.03}]
        ),
        "backtest_metrics": {"sharpe": 1.2, "deflated_sharpe": 0.8, "q5_q1_spread_net": 0.06},
    }
    with phase4_path.open("wb") as handle:
        pickle.dump(payload, handle)

    tables = _phase4_tables(_load_phase4_results(phase4_path))

    assert tables["IC results"][0]["variable"] == "sentiment_dynamics"
    assert tables["Fama-MacBeth coefficients"][0]["coef"] == 0.04
    assert {"metric": "sharpe", "value": 1.2} in tables["Backtest metrics"]


def test_report_builder_injects_phase4_tables_and_claude_markers(tmp_path: Path):
    phase4_path = tmp_path / "phase4_results.pkl"
    output_path = tmp_path / "report.docx"
    with phase4_path.open("wb") as handle:
        pickle.dump(
            {
                "ic_results": [{"variable": "governance_credibility", "ic_mean": 0.07, "t_nw": 1.9}],
                "fama_macbeth_coefficients": [{"variable": "governance_credibility", "coef": 0.02, "p": 0.05}],
                "backtest_metrics": {"sharpe": 1.1},
            },
            handle,
        )

    build_final_report(output_path, phase4_path=phase4_path)
    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)

    assert output_path.exists()
    assert "Phase 4 Validation Tables" in text
    assert CLAUDE_MARKER in table_text
    assert "IC results" in text
