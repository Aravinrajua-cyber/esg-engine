from __future__ import annotations

import json
import pickle
from pathlib import Path

import pandas as pd
from docx import Document
from pptx import Presentation

from tools.submit.assemble_submission import CLAUDE_MARKER, assemble_submission


PROJECT_ROOT = Path(__file__).resolve().parents[1]


def _write_json(path: Path, payload: dict) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload), encoding="utf-8")
    return path


def test_assembler_builds_submission_from_mock_phase4(tmp_path: Path):
    phase4_path = tmp_path / "phase4_results.pkl"
    with phase4_path.open("wb") as handle:
        pickle.dump(
            {
                "ic_results": pd.DataFrame(
                    [{"variable": "sentiment_dynamics", "horizon": 3, "ic_mean": 0.12, "t_nw": 2.1, "fdr_survived": True}]
                ),
                "fama_macbeth_coefficients": pd.DataFrame(
                    [{"variable": "sentiment_dynamics", "coef": 0.04, "t_nw": 2.4, "p": 0.03}]
                ),
                "backtest_metrics": {"sharpe": 1.2, "deflated_sharpe": 0.8, "q5_q1_spread_net": 0.06},
            },
            handle,
        )
    site_content_path = _write_json(tmp_path / "site_content.json", {"brand": "ESG Momentum Engine"})
    site_data_path = _write_json(tmp_path / "site_data.json", {"data_mode": "synthetic"})

    result = assemble_submission(
        root=PROJECT_ROOT,
        output_dir=tmp_path / "submission",
        phase4_path=phase4_path,
        site_content_path=site_content_path,
        site_data_path=site_data_path,
        run_site_build=False,
    )

    assert result.report_path.exists()
    assert result.deck_path.exists()
    assert result.readme_path.exists()
    assert result.site_path is None
    assert result.figure_count >= 1
    assert result.remaining_markers == 0

    document = Document(result.report_path)
    doc_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert CLAUDE_MARKER not in doc_text + table_text
    assert "sentiment_dynamics" in table_text

    deck = Presentation(result.deck_path)
    assert len(deck.slides) == 12

    readme = result.readme_path.read_text(encoding="utf-8")
    assert "ESG_Momentum_Engine_Report.docx" in readme
    assert "site/` - not generated" in readme
