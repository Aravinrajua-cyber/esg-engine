"""Build Word reports for the ESG Momentum Engine.

The final report is a synthetic-demo technical report. It uses generated JSON
artifacts and rendered PNG figures, and it labels all results as illustrative.
"""

from __future__ import annotations

import json
import math
import pickle
from collections import Counter
from datetime import date
from pathlib import Path
from statistics import mean
from typing import Any, Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = PROJECT_ROOT / "outputs" / "report"
FIGURE_DIR = PROJECT_ROOT / "outputs" / "figures"
SITE_DATA_DIR = PROJECT_ROOT / "outputs" / "site_data"
PHASE4_RESULTS_PATH = PROJECT_ROOT / "data" / "processed" / "phase4_results.pkl"

OUT_PATH = REPORT_DIR / "ESG_Momentum_Engine_Report.docx"
FINAL_OUT_PATH = REPORT_DIR / "ESG_Momentum_Engine_Final_Report.docx"

ACCENT = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 11, 12)
MUTED = RGBColor(85, 85, 85)
WHITE = RGBColor(255, 255, 255)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
RISK_FILL = "FFF2CC"
CLAUDE_MARKER = "[CLAUDE_WRITES_HERE]"

PILLARS = [
    "sentiment_dynamics",
    "transition_readiness",
    "governance_credibility",
    "disclosure_behavior",
]

FIGURES = [
    (
        "treemap.png",
        "Universe composition treemap",
        "This synthetic treemap shows how the demonstration universe is distributed across sectors and markets. It is useful for checking whether the demo feed is broad enough for frontend and report testing.",
    ),
    (
        "coverage_heatmap.png",
        "Coverage heatmap",
        "The heatmap highlights coverage variation by country and sector. In live mode, this chart should be used to identify where missing source data may weaken model confidence.",
    ),
    (
        "ic_bar.png",
        "Per-variable information coefficient bar chart",
        "This chart illustrates how validation evidence would be presented for individual variables. The current values are synthetic placeholders and should not be read as empirical proof.",
    ),
    (
        "ic_decay_curves.png",
        "Information-coefficient decay curves",
        "The decay curves show the intended diagnostic for whether a signal weakens over longer horizons. In synthetic mode, the shape demonstrates the visual contract only.",
    ),
    (
        "money_chart.png",
        "Money chart: Q5, Q1, benchmark, and naive ESG",
        "This chart is the primary performance display. It compares cumulative synthetic paths and includes a train/test boundary to show how live validation should be separated.",
    ),
    (
        "walk_forward_weight_ribbon.png",
        "Walk-forward weight ribbon",
        "The ribbon displays how composite weights can evolve through walk-forward validation. The demo version confirms report readability and chart plumbing.",
    ),
    (
        "placebo_histogram.png",
        "Placebo histogram",
        "The histogram demonstrates the intended placebo-test view: a realized spread compared against randomized outcomes. Synthetic values are illustrative only.",
    ),
    (
        "matrix_scatter.png",
        "ESG level versus ESG momentum matrix",
        "The scatter plot maps companies into the four classification quadrants. This is the report version of the company-detail 2x2 view used by the frontend.",
    ),
    (
        "equity_drawdown.png",
        "Equity curve and drawdown subplot",
        "The paired chart shows cumulative performance together with drawdown. It is designed to prevent performance interpretation without downside context.",
    ),
    (
        "by_country_sector_spreads.png",
        "By-country and by-sector spread bars",
        "The bar chart shows where synthetic spreads appear stronger or weaker. In live mode, this will be a key diagnostic for market-specific effects.",
    ),
    (
        "composite_correlation_heatmap.png",
        "Composite correlation heatmap",
        "This figure checks whether composite signals are too redundant. High live correlations would require tighter governance before deployment.",
    ),
    (
        "score_histogram.png",
        "Score histogram with grade bands",
        "The histogram explains the distribution of final scores and the placement of grade bands. It helps reviewers spot compression or unrealistic score clustering.",
    ),
    (
        "top20_leaderboard.png",
        "Top-20 leaderboard table figure",
        "The leaderboard figure gives a print-friendly snapshot of the highest-ranked synthetic companies. It must remain labelled as demonstration data.",
    ),
    (
        "company_overlay.png",
        "Per-company price, sentiment, and score overlay",
        "The overlay demonstrates the single-company diagnostic view. It is intended to show whether score movement is directionally consistent with source proxies.",
    ),
]

FLAG_WEIGHTS = {
    "LOW_COVERAGE": 30,
    "CONTROVERSY_RISING": 25,
    "LOW_LIQUIDITY": 20,
    "HIGH_VOL": 15,
    "STALE_DATA": 10,
}


def _load_json(name: str) -> Any:
    return json.loads((SITE_DATA_DIR / name).read_text(encoding="utf-8"))


def _load_phase4_results(path: Path = PHASE4_RESULTS_PATH) -> dict[str, Any]:
    """Load Claude-owned Phase 4 validation outputs when available."""

    if not path.exists():
        return {}
    with path.open("rb") as handle:
        data = pickle.load(handle)
    if not isinstance(data, dict):
        raise TypeError("phase4_results.pkl must contain a dict-like results object")
    return data


def _fmt(value: Any, digits: int = 1) -> str:
    if value is None:
        return "n/a"
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, int):
        return f"{value:,}"
    if isinstance(value, float):
        if math.isnan(value):
            return "n/a"
        return f"{value:,.{digits}f}"
    return str(value)


def _is_missing(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, float) and math.isnan(value):
        return True
    return False


def _rows_from_table(value: Any) -> list[dict[str, Any]]:
    if value is None:
        return []
    if isinstance(value, pd.DataFrame):
        return value.reset_index(drop=False).to_dict("records")
    if isinstance(value, pd.Series):
        return [{"metric": key, "value": val} for key, val in value.to_dict().items()]
    if isinstance(value, list):
        rows = []
        for item in value:
            if isinstance(item, dict):
                rows.append(item)
            else:
                rows.append({"value": item})
        return rows
    if isinstance(value, dict):
        if all(not isinstance(item, (dict, list, tuple, pd.DataFrame, pd.Series)) for item in value.values()):
            return [{"metric": key, "value": val} for key, val in value.items()]
        rows = []
        for key, item in value.items():
            if isinstance(item, dict):
                rows.append({"name": key, **item})
            else:
                rows.append({"name": key, "value": item})
        return rows
    return [{"value": value}]


def _first_table(results: dict[str, Any], keys: Iterable[str]) -> list[dict[str, Any]]:
    for key in keys:
        if key in results:
            rows = _rows_from_table(results[key])
            if rows:
                return rows
    return []


def _phase4_tables(results: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    return {
        "IC results": _first_table(results, ("ic_results", "ic", "ic_table", "information_coefficients")),
        "Fama-MacBeth coefficients": _first_table(
            results,
            ("fama_macbeth_coefficients", "fama_macbeth", "fama_macbeth_results", "fm_results"),
        ),
        "Backtest metrics": _first_table(results, ("backtest_metrics", "backtest", "performance_metrics")),
    }


def _ordered_columns(rows: list[dict[str, Any]], preferred: Iterable[str]) -> list[str]:
    present = {column for row in rows for column in row.keys()}
    ordered = [column for column in preferred if column in present]
    ordered.extend(sorted(present - set(ordered)))
    return ordered


def _add_claude_marker(doc: Document, section: str) -> None:
    _callout(
        doc,
        CLAUDE_MARKER,
        f"Claude will add final {section} narrative here. Codex leaves this marker intentionally.",
        RISK_FILL,
    )


def _add_dataframe_like_table(
    doc: Document,
    title: str,
    rows: list[dict[str, Any]],
    preferred_columns: Iterable[str],
) -> None:
    doc.add_heading(title, level=2)
    if not rows:
        _callout(doc, "Phase 4 table unavailable.", f"{CLAUDE_MARKER} Claude will populate {title}.", RISK_FILL)
        return
    columns = _ordered_columns(rows, preferred_columns)
    widths = [6.5 / max(1, len(columns))] * len(columns)
    _table(doc, columns, ([row.get(column) for column in columns] for row in rows), widths)


def _risk_index(company: dict[str, Any]) -> int:
    missing_component = max(0, 100 - float(company.get("coverage_pct", 0)))
    flag_component = sum(FLAG_WEIGHTS.get(flag, 0) for flag in company.get("flags", []))
    return int(round(min(100, missing_component + flag_component)))


def _stats(companies: list[dict[str, Any]]) -> dict[str, Any]:
    scores = [float(c["overall_score"]) for c in companies]
    coverage = [float(c["coverage_pct"]) for c in companies]
    risks = [_risk_index(c) for c in companies]
    return {
        "count": len(companies),
        "score_min": min(scores),
        "score_mean": mean(scores),
        "score_max": max(scores),
        "coverage_mean": mean(coverage),
        "risk_mean": mean(risks),
        "timeseries_count": sum(1 for c in companies if c.get("timeseries") is not None),
        "grade_counts": Counter(c["grade"] for c in companies),
        "classification_counts": Counter(c["classification"] for c in companies),
        "country_counts": Counter(c["country"] for c in companies),
    }


def _set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, end):
        run._r.append(element)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        _shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _set_run_font(run, 9.5, INK, bold=bold)


def _set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def _table(doc: Document, headers: list[str], rows: Iterable[Iterable[Any]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for cell, width in zip(table.rows[0].cells, widths):
            _set_cell_width(cell, width)
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, bold=True, fill=HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        if widths:
            for cell, width in zip(cells, widths):
                _set_cell_width(cell, width)
        for cell, value in zip(cells, row):
            _set_cell_text(cell, _fmt(value))
    doc.add_paragraph()


def _paragraph(doc: Document, text: str, style: str | None = None, bold_label: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_label:
        run = p.add_run(bold_label)
        _set_run_font(run, bold=True)
        text = text.removeprefix(bold_label)
    run = p.add_run(text)
    _set_run_font(run, 11, INK)


def _bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, 11, INK)


def _numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        _set_run_font(run, 11, INK)


def _callout(doc: Document, label: str, body: str, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.rows[0].cells[0]
    _set_cell_width(cell, 6.5)
    _shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + " ")
    _set_run_font(r1, 10.5, DARK_BLUE, bold=True)
    r2 = p.add_run(body)
    _set_run_font(r2, 10.5, INK)
    doc.add_paragraph()


def _add_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    _add_field(p, r'TOC \o "1-3" \h \z \u')
    note = doc.add_paragraph()
    run = note.add_run("In Word, right-click this table and choose Update Field to refresh page numbers.")
    _set_run_font(run, 9, MUTED, italic=True)
    doc.add_page_break()


def _add_figure(doc: Document, index: int, filename: str, caption: str, interpretation: str) -> None:
    path = FIGURE_DIR / filename
    doc.add_heading(f"Figure {index}. {caption}", level=2)
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.35))
        picture = doc.paragraphs[-1]
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        _callout(doc, "Missing figure.", f"{path} was not found.", RISK_FILL)
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(f"Figure {index}. {caption}. Source: generated synthetic-demo artifact {filename}.")
    _paragraph(doc, interpretation, bold_label="")


def _apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after, color in [
        ("Title", 24, 0, 6, INK),
        ("Subtitle", 12, 0, 14, MUTED),
        ("Heading 1", 16, 16, 8, ACCENT),
        ("Heading 2", 13, 12, 6, ACCENT),
        ("Heading 3", 12, 8, 4, DARK_BLUE),
        ("Caption", 9, 3, 8, MUTED),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def _set_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    run = header.add_run("ESG Momentum Engine | Synthetic-demo report")
    _set_run_font(run, 9, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Page ")
    _set_run_font(r, 9, MUTED)
    _add_field(footer, "PAGE")


def _cover(doc: Document, feed: dict[str, Any], stats: dict[str, Any]) -> None:
    title = doc.add_paragraph(style="Title")
    title.add_run("ESG Momentum Engine")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Final synthetic-demo research report and implementation handoff")
    rows = [
        ("Prepared for", "PolyFinTech100 2026 / CGS International ESG Intelligence"),
        ("Generated", date.today().isoformat()),
        ("Data mode", feed["data_mode"].upper()),
        ("Universe size", stats["count"]),
        ("Primary artifact", "outputs/site_data/companies.json"),
        ("Report status", "Synthetic demonstration. Not investment advice."),
    ]
    _table(doc, ["Field", "Value"], rows, [1.75, 4.75])
    _callout(
        doc,
        "Important.",
        "Every score, spread, validation statistic, ranking, and chart in this report is illustrative synthetic-demo output unless explicitly marked otherwise. The report validates product plumbing and methodology presentation, not investable findings.",
        RISK_FILL,
    )
    doc.add_page_break()


def _section_intro(doc: Document, title: str, body: str) -> None:
    doc.add_heading(title, level=1)
    _paragraph(doc, body)


def build_final_report(out_path: Path = FINAL_OUT_PATH, phase4_path: Path = PHASE4_RESULTS_PATH) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    feed = _load_json("companies.json")
    backtest = _load_json("backtest.json")
    placebo = _load_json("placebo.json")
    ic_rows = _load_json("ic_table.json")
    phase4_results = _load_phase4_results(phase4_path)
    phase4_tables = _phase4_tables(phase4_results)
    companies = feed["companies"]
    stats = _stats(companies)

    doc = Document()
    _apply_styles(doc)
    _set_header_footer(doc)
    _cover(doc, feed, stats)
    _add_toc(doc)

    _section_intro(
        doc,
        "1. Executive Summary",
        "The ESG Momentum Engine is a research and product prototype for identifying companies whose ESG trajectory may be improving before slower static ratings fully reflect that change. This final report documents the synthetic demonstration build, the scoring contract, the frontend-ready data artifacts, the figure suite, and the governance limits that must be satisfied before live use.",
    )
    _bullets(
        doc,
        [
            f"The current feed contains {stats['count']} synthetic company records and is explicitly labelled data_mode=synthetic.",
            f"The mean synthetic overall score is {_fmt(stats['score_mean'])}, with a range from {_fmt(stats['score_min'])} to {_fmt(stats['score_max'])}.",
            f"{stats['timeseries_count']} companies include synthetic timeseries overlays; the rest intentionally use null timeseries values to exercise missing-data states.",
            "No output in this report should be treated as a live company recommendation, valuation signal, or investment advice.",
        ],
    )

    _section_intro(
        doc,
        "2. Problem Statement",
        "Conventional ESG ratings can be slow-moving and difficult to connect to market timing. The prototype asks whether faster-moving indicators such as sentiment dynamics, transition-readiness proxies, governance credibility, and disclosure behavior can be composed into a transparent momentum score for research workflows.",
    )

    _section_intro(
        doc,
        "3. Intended User Segment",
        "The intended users are research analysts, ESG product specialists, model governance reviewers, and institutional-client teams who need an explainable ranking interface rather than a black-box score. The design emphasizes traceability, caveats, and drill-down review before any downstream action.",
    )

    _section_intro(
        doc,
        "4. ASEAN-Focused Scope",
        "The project scope is ASEAN-focused. The configured markets include Singapore, Indonesia, Malaysia, Thailand, the Philippines, and Vietnam. The synthetic feed preserves country, exchange, sector, currency, market-cap tier, and liquidity-style fields so the frontend and reporting layers can be tested against the intended regional structure.",
    )

    _section_intro(
        doc,
        "5. Real-Data Mode Versus Synthetic-Demo Mode",
        "Real-data mode will depend on live fetchers and validation artifacts. Synthetic-demo mode is a deterministic seeded generator used to prove the schema, UI, charts, and reporting workflow while avoiding fabricated claims about real companies.",
    )
    _table(
        doc,
        ["Mode", "Purpose", "Permitted interpretation"],
        [
            ("Synthetic demo", "Exercise schema, UX, charts, tests, and report layout.", "Illustrative only; no investable conclusion."),
            ("Live data", "Use fetched prices, FX, fundamentals, ESG snapshots, disclosures, and sentiment.", "Research output subject to validation, licensing, and governance review."),
        ],
        [1.4, 2.7, 2.4],
    )

    _section_intro(
        doc,
        "6. Data Schema",
        "The frozen schema is maintained in SCHEMAS.md. The primary frontend feed is companies.json with top-level metadata, model weights, pillar definitions, risk-flag definitions, and a list of company records.",
    )
    _table(
        doc,
        ["Artifact", "Role", "Current status"],
        [
            ("companies.json", "Primary company ranking feed.", f"{stats['count']} synthetic records generated."),
            ("backtest.json", "Q5, Q1, benchmark, and naive ESG paths.", f"{len(backtest['dates'])} synthetic dates."),
            ("ic_table.json", "Per-variable validation rows for methodology charts.", f"{len(ic_rows)} synthetic rows."),
            ("placebo.json", "Placebo histogram and realized-spread marker.", f"{len(placebo['hist_bins'])} bins."),
            ("by_country.json / by_sector.json", "Grouped spread charts.", "Synthetic aggregates generated."),
        ],
        [1.9, 3.1, 1.5],
    )

    _section_intro(
        doc,
        "7. Synthetic Data-Generation Methodology",
        "The synthetic generator uses the configured random seed 42 and the available universe parquet as its source shape. If fewer than 500 source companies are available, it pads with clearly named synthetic demo companies. Scores are generated from bounded random pillar values, then ranked deterministically by overall score.",
    )
    _add_claude_marker(doc, "methodology")
    _numbered(
        doc,
        [
            "Read universe fields such as ticker, name, country, exchange, sector, tier, currency, and liquidity flag.",
            "Generate four weighted pillar scores and a display-only data coverage score on a 0 to 100 scale.",
            "Compute an overall synthetic score from the weighted pillar sum with small seeded noise.",
            "Assign confidence bands, classifications, flags, and top-company synthetic timeseries.",
            "Sort by overall score and assign ranks from 1 to 500.",
        ],
    )

    _section_intro(
        doc,
        "8. Scoring Methodology",
        "The validated weight vector in the synthetic feed is applied to four model pillars. Data coverage is displayed to the user and used in risk/confidence interpretation, but it is excluded from the weighted score calculation.",
    )
    _add_claude_marker(doc, "scoring methodology")
    weights = feed["model"]["validated_weights"]
    _table(
        doc,
        ["Pillar", "Weight", "Interpretation"],
        [
            ("Sentiment dynamics", weights["sentiment_dynamics"], "Direction and pace of news-tone and attention proxies."),
            ("Transition readiness", weights["transition_readiness"], "Capital-allocation and transition-preparedness proxies."),
            ("Governance credibility", weights["governance_credibility"], "Governance quality, consistency, and credibility proxies."),
            ("Disclosure behavior", weights["disclosure_behavior"], "Frequency and quality of sustainability-related disclosure behavior."),
        ],
        [2.0, 1.0, 3.5],
    )

    _section_intro(
        doc,
        "9. Full Explanation Of Every Score Component",
        "Each component is designed to be interpretable on its own before it contributes to the final score. In synthetic mode these are plausible placeholders; in live mode each must trace to an auditable source field and calculation.",
    )
    _table(
        doc,
        ["Component", "Bounds", "Used in overall score", "Plain-language meaning"],
        [
            ("sentiment_dynamics", "0-100", "Yes", "Whether recent ESG-related attention and tone are improving."),
            ("transition_readiness", "0-100", "Yes", "Whether capital allocation and financial proxies suggest transition preparedness."),
            ("governance_credibility", "0-100", "Yes", "Whether governance behavior looks credible and consistent."),
            ("disclosure_behavior", "0-100", "Yes", "Whether sustainability disclosure behavior is timely and persistent."),
            ("data_coverage", "0-100", "No", "How complete the expected input coverage is for the company."),
        ],
        [1.65, 0.8, 1.25, 2.8],
    )

    _section_intro(
        doc,
        "10. Overall-Score Calculation",
        "The frontend weight sandbox recomputes the same score formula client-side. The display-only data_coverage pillar is excluded from the sum so users can see coverage without letting it masquerade as alpha.",
    )
    _callout(
        doc,
        "Formula.",
        "overall_score = 0.35 * sentiment_dynamics + 0.25 * transition_readiness + 0.25 * governance_credibility + 0.15 * disclosure_behavior. The synthetic generator then adds seeded demonstration noise before ranking.",
    )

    _section_intro(
        doc,
        "11. Phase 4 Validation Tables",
        "This section reads Claude-owned validation output from data/processed/phase4_results.pkl. When that pickle is absent, Codex leaves explicit placeholders instead of inventing statistics.",
    )
    _add_dataframe_like_table(
        doc,
        "IC results",
        phase4_tables["IC results"],
        ("variable", "label", "horizon", "ic_mean", "ic_std", "ic_3m", "t_nw", "hit_rate", "p", "fdr_survived"),
    )
    _add_dataframe_like_table(
        doc,
        "Fama-MacBeth coefficients",
        phase4_tables["Fama-MacBeth coefficients"],
        ("variable", "coef", "t_nw", "p", "std_err", "fdr_survived"),
    )
    _add_dataframe_like_table(
        doc,
        "Backtest metrics",
        phase4_tables["Backtest metrics"],
        ("metric", "value", "winning_composite", "gross", "net", "sharpe", "sortino", "max_dd", "calmar", "q5_q1_spread_net", "deflated_sharpe"),
    )

    _section_intro(
        doc,
        "12. Risk-Index Calculation",
        "The frozen companies.json schema stores risk flags and coverage, not a separate risk_index field. For reporting and QA, a derived risk index is calculated from missing coverage plus weighted flags and clipped to a 0 to 100 range.",
    )
    _table(
        doc,
        ["Risk input", "Weight or role"],
        [
            ("Missing coverage component", "100 - coverage_pct"),
            ("LOW_COVERAGE", FLAG_WEIGHTS["LOW_COVERAGE"]),
            ("CONTROVERSY_RISING", FLAG_WEIGHTS["CONTROVERSY_RISING"]),
            ("LOW_LIQUIDITY", FLAG_WEIGHTS["LOW_LIQUIDITY"]),
            ("HIGH_VOL", FLAG_WEIGHTS["HIGH_VOL"]),
            ("STALE_DATA", FLAG_WEIGHTS["STALE_DATA"]),
        ],
        [3.25, 3.25],
    )
    _callout(doc, "Formula.", "risk_index = min(100, max(0, 100 - coverage_pct) + sum(flag weights)).", RISK_FILL)

    _section_intro(
        doc,
        "13. Confidence-Band Methodology",
        "Confidence bands represent uncertainty around the score. In synthetic mode, the band width is generated from a seeded range and clipped to 0 to 100. In live mode, it should be replaced by a statistically grounded uncertainty estimate that reflects input coverage, source freshness, validation error, and market-specific reliability.",
    )
    _add_claude_marker(doc, "confidence methodology")

    _section_intro(
        doc,
        "14. Descriptive Statistics",
        "The following statistics summarize the generated synthetic feed. They are included to verify distribution sanity and report plumbing, not to describe live ASEAN issuers.",
    )
    _table(
        doc,
        ["Metric", "Value"],
        [
            ("Company records", stats["count"]),
            ("Mean overall score", stats["score_mean"]),
            ("Minimum overall score", stats["score_min"]),
            ("Maximum overall score", stats["score_max"]),
            ("Mean coverage", stats["coverage_mean"]),
            ("Mean derived risk index", stats["risk_mean"]),
            ("Companies with timeseries", stats["timeseries_count"]),
        ],
        [3.25, 3.25],
    )
    _table(doc, ["Grade", "Count"], sorted(stats["grade_counts"].items()), [3.25, 3.25])
    _table(doc, ["Classification", "Count"], sorted(stats["classification_counts"].items()), [3.25, 3.25])

    doc.add_heading("15. Figures And Plain-Language Interpretations", level=1)
    _paragraph(
        doc,
        "All figures below are generated from synthetic or fixture-style artifacts. They are inserted at near full page width so they can be read in Word without excessive zooming.",
    )
    for i, (filename, caption, interpretation) in enumerate(FIGURES, 1):
        _add_figure(doc, i, filename, caption, interpretation)
        if i in {4, 8, 11}:
            doc.add_page_break()

    _section_intro(
        doc,
        "16. Limitations",
        "The current build validates data contracts, visualization plumbing, and report presentation. It does not validate a live investment model. Synthetic scores, spreads, ICs, and rankings should not be used outside demonstration and testing contexts.",
    )
    _add_claude_marker(doc, "limitations and risk prose")
    _bullets(
        doc,
        [
            "Synthetic values can look realistic while containing no empirical evidence.",
            "Live fetchers may have exchange-specific gaps, stale fields, and licensing constraints.",
            "Statistical validation is not complete until Claude-owned live validation artifacts are produced.",
            "Frontend visual QA is blocked until npm dependencies can be installed.",
        ],
    )

    _section_intro(
        doc,
        "17. Risk Register",
        "The risk register below records the main implementation and interpretation risks visible at this stage.",
    )
    _add_claude_marker(doc, "risk register narrative")
    _table(
        doc,
        ["Risk", "Severity", "Mitigation"],
        [
            ("Synthetic output mistaken for live findings.", "High", "Persistent synthetic banner, report disclaimers, tests for data_mode=synthetic."),
            ("Incomplete source coverage in live mode.", "High", "Coverage metrics, confidence bands, failure logs, missing values kept as null/NaN."),
            ("Overfitted signal design.", "High", "Train/test split, placebo tests, walk-forward validation, FDR review."),
            ("Frontend dependency installation blocked.", "Medium", "Document blocker and rerun npm commands in normal terminal."),
            ("Unclear licensing for downstream display.", "High", "Review data-source terms before live distribution."),
        ],
        [2.0, 1.0, 3.5],
    )

    _section_intro(
        doc,
        "18. Data-Source And Licensing Risks",
        "Live mode may use yfinance/Yahoo Finance-derived market data, GDELT news data, exchange announcements, and issuer disclosures. Each source must be reviewed for permitted use, redistribution, caching, attribution, and client-facing display rights before launch.",
    )
    _add_claude_marker(doc, "data-source and licensing risks")

    _section_intro(
        doc,
        "19. Statistical Risks",
        "Statistical risks include look-ahead bias, survivorship bias, multiple testing, factor crowding, small sample sizes by market, unstable correlations, non-stationary relationships, transaction-cost underestimation, and overstated inference from noisy proxies.",
    )
    _add_claude_marker(doc, "statistical risks")

    _section_intro(
        doc,
        "20. Model-Governance Risks",
        "The model requires versioned data contracts, frozen validation windows, documented parameter changes, reviewable transformations, and explicit sign-off before live client use. Weight sandbox outputs should remain labelled as custom exploratory views rather than validated model outputs.",
    )
    _add_claude_marker(doc, "model-governance risks")

    _section_intro(
        doc,
        "21. Cybersecurity And Privacy Risks",
        "The current frontend is static and reads JSON only. Future upload, API, or CSV-import paths must validate file type, size, schema, encoding, numeric bounds, formula injection risks, and personally identifiable information before accepting user-provided data.",
    )
    _add_claude_marker(doc, "cybersecurity and privacy risks")

    _section_intro(
        doc,
        "22. User-Misinterpretation Risks",
        "Users may overread ranks, grade badges, confidence bands, or backtest-style charts. The product should keep disclaimers visible, explain missing data, show confidence and coverage next to scores, and avoid action-oriented language such as buy, sell, or outperform.",
    )
    _add_claude_marker(doc, "user-misinterpretation risks")

    _section_intro(
        doc,
        "23. Accessibility Considerations",
        "The frontend source includes semantic sections, keyboard-open behavior for leaderboard rows, responsive mobile collapse, persistent labels, and reduced-motion CSS. Browser-based accessibility validation is still required once npm dependencies can be installed and the Vite app can run.",
    )

    _section_intro(
        doc,
        "24. Future Real-Data Integration Plan",
        "The next phase should connect the existing fetchers to live artifacts, backfill price/FX/fundamental/ESG/disclosure/sentiment data, generate Claude-owned validation results, and replace synthetic frontend feeds only after schema and governance checks pass.",
    )
    _numbered(
        doc,
        [
            "Run fetchers with cache and failure logs enabled.",
            "Validate raw and processed parquet schemas against SCHEMAS.md.",
            "Generate signal_panel, returns_forward, and validation_results artifacts.",
            "Regenerate site_data in live mode with retrieval dates and coverage metrics.",
            "Run pytest, build the frontend, and complete browser visual QA.",
            "Review licensing and governance sign-offs before external use.",
        ],
    )

    _section_intro(
        doc,
        "25. Reproducibility Instructions",
        "The following commands reproduce the current non-frontend artifacts in this environment. Frontend commands require npm dependency access as documented in LOCAL_RUN_GUIDE.md.",
    )
    _table(
        doc,
        ["Task", "Command"],
        [
            ("Generate synthetic site data", r".\.venv\Scripts\python.exe -m src.fetchers.synthetic"),
            ("Render figures", r".\.venv\Scripts\python.exe -m src.viz.style"),
            ("Build final report", r".\.venv\Scripts\python.exe -m src.report.build_report"),
            ("Run tests", r".\.venv\Scripts\python.exe -m pytest -q"),
            ("Frontend install, normal terminal", r"cd /d C:\Hackathon\esg-engine\site && npm.cmd install"),
            ("Frontend build, normal terminal", r"cd /d C:\Hackathon\esg-engine\site && npm.cmd run build"),
        ],
        [2.25, 4.25],
    )

    _section_intro(
        doc,
        "26. Bibliography With Citations",
        "These references identify the source systems, libraries, and artifacts relevant to the demonstration. They are not evidence for the synthetic numerical results.",
    )
    _bullets(
        doc,
        [
            "Project contract: SCHEMAS.md, schema_version 1, local repository artifact.",
            "Project configuration: config/settings.yaml, local repository artifact.",
            "Yahoo Finance data access layer: yfinance Python package documentation and Yahoo Finance public data surfaces.",
            "GDELT Project: GDELT 2.0 DOC API documentation for news tone and volume concepts.",
            "Singapore Exchange and Bursa Malaysia announcement pages for future disclosure-source integration.",
            "Python libraries: pandas, numpy, plotly, python-docx, pytest.",
            "Frontend libraries: React, Vite, TypeScript, Plotly.js, lucide-react.",
        ],
    )

    _section_intro(
        doc,
        "27. Appendix A: Full Variable Dictionary",
        "The dictionary below covers the primary companies.json feed and the derived reporting-only risk index.",
    )
    dictionary_rows = [
        ("schema_version", "Top-level", "Contract version. Current value is 1."),
        ("generated_at", "Top-level", "UTC generation timestamp."),
        ("data_mode", "Top-level", "live or synthetic; current report requires synthetic."),
        ("as_of_date", "Top-level", "Date represented by the generated feed."),
        ("universe_size", "Top-level", "Number of companies in the feed."),
        ("winning_composite", "Model", "Name of the selected composite."),
        ("validated_weights", "Model", "Weights used by the default score calculation."),
        ("train_end", "Model", "End date for training window in live validation."),
        ("headline", "Model", "Headline model metrics for the hero and summary UI."),
        ("ticker", "Company", "Ticker symbol."),
        ("name", "Company", "Issuer or synthetic demo company name."),
        ("country", "Company", "Market country code or name from the universe shape."),
        ("exchange", "Company", "Exchange identifier."),
        ("sector", "Company", "Sector grouping."),
        ("mcap_tier", "Company", "mega, large, or mid."),
        ("currency", "Company", "Local trading/reporting currency."),
        ("rank", "Company", "Rank after sorting by overall_score descending."),
        ("overall_score", "Company", "Final 0 to 100 model score."),
        ("grade", "Company", "A+, A, B, C, or D score band."),
        ("confidence_low", "Company", "Lower bound of score confidence interval."),
        ("confidence_high", "Company", "Upper bound of score confidence interval."),
        ("coverage_pct", "Company", "Expected-input coverage percentage."),
        ("classification", "Company", "One of hidden_winner, future_leader, overrated_leader, value_trap."),
        ("esg_level_pctile", "Company", "0 to 100 axis for static ESG level."),
        ("esg_momentum_pctile", "Company", "0 to 100 axis for ESG momentum."),
        ("pillar_scores", "Company", "Five component scores: four weighted pillars plus display-only data coverage."),
        ("flags", "Company", "Risk flags from the schema-defined flag list."),
        ("explanation", "Company", "Deterministic explanation sentence."),
        ("timeseries", "Company", "Null or arrays for dates, price_usd, sentiment_tone, and score."),
        ("risk_index", "Derived report metric", "Not stored in companies.json. Derived as min(100, 100 - coverage_pct + weighted flags)."),
    ]
    _table(doc, ["Variable", "Level", "Definition"], dictionary_rows, [1.7, 1.25, 3.55])

    _section_intro(
        doc,
        "28. Appendix B: Generated Artifact Inventory",
        "The final report was generated from the local artifacts listed below.",
    )
    _table(
        doc,
        ["Artifact", "Path"],
        [
            ("Companies feed", str(SITE_DATA_DIR / "companies.json")),
            ("Backtest feed", str(SITE_DATA_DIR / "backtest.json")),
            ("IC table", str(SITE_DATA_DIR / "ic_table.json")),
            ("Placebo feed", str(SITE_DATA_DIR / "placebo.json")),
            ("Phase 4 results", str(phase4_path)),
            ("Figures directory", str(FIGURE_DIR)),
            ("Local run guide", str(PROJECT_ROOT / "LOCAL_RUN_GUIDE.md")),
        ],
        [2.0, 4.5],
    )

    doc.save(out_path)
    return out_path


def build_report(out_path: Path = OUT_PATH, phase4_path: Path = PHASE4_RESULTS_PATH) -> Path:
    """Backward-compatible report builder entry point."""

    return build_final_report(out_path, phase4_path)


if __name__ == "__main__":
    print(build_final_report())
