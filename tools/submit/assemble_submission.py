from __future__ import annotations

import argparse
import json
import pickle
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import yaml
from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))  # allow `python tools/submit/assemble_submission.py`
SUBMISSION_DIR = PROJECT_ROOT / "outputs" / "submission"
PHASE4_PATH = PROJECT_ROOT / "data" / "processed" / "phase4_results.pkl"
FIGURES_DIR = PROJECT_ROOT / "outputs" / "figures"
SITE_DIR = PROJECT_ROOT / "site"
SITE_CONTENT_PATH = SITE_DIR / "src" / "site_content.json"
# The frontend data contract is outputs/site_data/*.json (companies.json carries data_mode);
# there is no site/src/site_data.json. README metadata reads data_mode from companies.json.
SITE_DATA_PATH = PROJECT_ROOT / "outputs" / "site_data" / "companies.json"
DECK_CONTENT_PATH = PROJECT_ROOT / "tools" / "deck" / "deck_content.yaml"
PROSE_PATH = PROJECT_ROOT / "tools" / "submit" / "submission_prose.yaml"
CLAUDE_MARKER = "[CLAUDE_WRITES_HERE]"


@dataclass
class SubmissionResult:
    output_dir: Path
    report_path: Path
    deck_path: Path
    site_path: Path | None
    readme_path: Path
    remaining_markers: int
    figure_count: int
    site_build_ran: bool


def _load_json(path: Path) -> Any:
    if not path.exists():
        raise FileNotFoundError(f"Required JSON input missing: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def _load_yaml(path: Path) -> dict[str, str]:
    if not path.exists():
        raise FileNotFoundError(f"Required YAML input missing: {path}")
    raw = yaml.safe_load(path.read_text(encoding="utf-8"))
    if not isinstance(raw, dict):
        raise ValueError(f"{path} must contain a mapping")
    return {str(key): str(value) for key, value in raw.items()}


def _load_phase4(path: Path) -> dict[str, Any]:
    if not path.exists():
        raise FileNotFoundError(f"Phase 4 results missing: {path}")
    with path.open("rb") as handle:
        payload = pickle.load(handle)
    if not isinstance(payload, dict):
        raise TypeError("phase4_results.pkl must contain a dict")
    return payload


def _reset_dir(path: Path) -> None:
    resolved = path.resolve()
    if resolved.exists():
        if resolved == resolved.anchor or "outputs" not in resolved.parts:
            raise ValueError(f"refusing to remove unsafe path: {resolved}")
        shutil.rmtree(resolved)
    resolved.mkdir(parents=True, exist_ok=True)


def _patch_report_builder_paths(root: Path, phase4_path: Path) -> None:
    from src.report import build_report as report_builder

    report_builder.PROJECT_ROOT = root
    report_builder.REPORT_DIR = root / "outputs" / "report"
    report_builder.FIGURE_DIR = root / "outputs" / "figures"
    report_builder.SITE_DATA_DIR = root / "outputs" / "site_data"
    report_builder.PHASE4_RESULTS_PATH = phase4_path


def _build_report(root: Path, output_path: Path, phase4_path: Path, prose_path: Path) -> tuple[Path, int]:
    from src.report.build_report import build_final_report

    _patch_report_builder_paths(root, phase4_path)
    build_final_report(output_path, phase4_path=phase4_path)
    prose = _load_yaml(prose_path)
    remaining = _fill_claude_markers(output_path, prose)
    return output_path, remaining


def _extract_prose_key(text: str) -> str | None:
    match = re.search(r"Claude will add final (.*?) narrative here", text)
    if match:
        return match.group(1).strip()
    match = re.search(r"Claude will populate (.*?)\.", text)
    if match:
        return match.group(1).strip()
    return None


def _prose_segments(text: str) -> list[tuple[bool, str]]:
    """Parse prose into (is_heading, text) segments. A line starting with '## ' is a bold
    subheading on its own paragraph; blank-line-separated runs of other lines are body
    paragraphs. Robust to '## h\\nbody' and '## h\\n\\nbody' alike."""
    segments: list[tuple[bool, str]] = []
    buf: list[str] = []

    def flush() -> None:
        joined = " ".join(line.strip() for line in buf if line.strip())
        if joined:
            segments.append((False, joined))
        buf.clear()

    for line in text.split("\n"):
        if line.startswith("## "):
            flush()
            segments.append((True, line[3:].strip()))
        elif not line.strip():
            flush()
        else:
            buf.append(line)
    flush()
    return segments or [(False, "")]


def _render_segment(paragraph, segment: tuple[bool, str]) -> None:
    is_heading, text = segment
    run = paragraph.add_run(text)
    run.bold = is_heading


def _replace_paragraph_text(paragraph, text: str) -> None:
    segments = _prose_segments(text)
    paragraph.clear()
    _render_segment(paragraph, segments[0])
    anchor = paragraph
    for segment in segments[1:]:
        new_p = anchor._p.makeelement(anchor._p.tag, {})
        anchor._p.addnext(new_p)
        following = type(paragraph)(new_p, paragraph._parent)
        _render_segment(following, segment)
        anchor = following


def _replace_cell_text(cell, text: str) -> None:
    segments = _prose_segments(text)
    for stale in cell.paragraphs[1:]:
        stale._element.getparent().remove(stale._element)
    first = cell.paragraphs[0]
    first.clear()
    _render_segment(first, segments[0])
    for segment in segments[1:]:
        _render_segment(cell.add_paragraph(), segment)


def _fill_claude_markers(docx_path: Path, prose: dict[str, str]) -> int:
    document = Document(docx_path)

    def replacement_for(text: str) -> str | None:
        if CLAUDE_MARKER not in text:
            return None
        key = _extract_prose_key(text)
        if key and key in prose:
            return prose[key]
        return prose.get("default")

    for paragraph in document.paragraphs:
        replacement = replacement_for(paragraph.text)
        if replacement:
            _replace_paragraph_text(paragraph, replacement)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                replacement = replacement_for(cell.text)
                if replacement:
                    _replace_cell_text(cell, replacement)

    document.save(docx_path)
    checked = Document(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in checked.paragraphs)
    table_text = "\n".join(cell.text for table in checked.tables for row in table.rows for cell in row.cells)
    return (paragraph_text + table_text).count(CLAUDE_MARKER)


def _run_site_build(site_dir: Path, output_site_dir: Path) -> Path:
    subprocess.run(["npm.cmd", "run", "build"], cwd=site_dir, check=True)
    dist_dir = site_dir / "dist"
    if not dist_dir.exists():
        raise FileNotFoundError(f"site build did not produce {dist_dir}")
    _reset_dir(output_site_dir)
    shutil.copytree(dist_dir, output_site_dir, dirs_exist_ok=True)
    return output_site_dir


def _build_deck(deck_content_path: Path, output_path: Path) -> Path:
    from tools.deck.build_deck import build_deck

    return build_deck(deck_content_path, output_path)


def _write_package_readme(
    output_dir: Path,
    result: SubmissionResult,
    site_content: dict[str, Any],
    site_data: dict[str, Any],
    phase4: dict[str, Any],
) -> Path:
    data_mode = site_data.get("data_mode", site_data.get("mode", "unknown"))
    lines = [
        "# ESG Momentum Engine Submission Package",
        "",
        "This folder contains the final generated submission artifacts.",
        "",
        "## Contents",
        "",
        f"- `ESG_Momentum_Engine_Report.docx` - final Word report assembled from Phase 4 results and figure PNGs.",
        f"- `ESG_Momentum_Engine.pptx` - final slide deck generated from `tools/deck/deck_content.yaml`.",
        f"- `site/` - static frontend export." if result.site_path else "- `site/` - not generated in this run (`--skip-site-build`).",
        "",
        "## Inputs Read",
        "",
        f"- Phase 4 keys: {', '.join(sorted(phase4.keys())) or 'none'}",
        f"- Figure PNG count: {result.figure_count}",
        f"- Site content brand: {site_content.get('brand', 'unknown')}",
        f"- Site data mode: {data_mode}",
        f"- Remaining Claude prose markers in report: {result.remaining_markers}",
        "",
        "## How To Rebuild",
        "",
        "```bat",
        "python tools/submit/assemble_submission.py",
        "```",
        "",
        "The build requires npm dependencies to be installed in `site/` for the final site export.",
        "",
    ]
    readme_path = output_dir / "README.md"
    readme_path.write_text("\n".join(lines), encoding="utf-8")
    return readme_path


def assemble_submission(
    *,
    root: Path = PROJECT_ROOT,
    output_dir: Path = SUBMISSION_DIR,
    phase4_path: Path = PHASE4_PATH,
    figures_dir: Path = FIGURES_DIR,
    site_content_path: Path = SITE_CONTENT_PATH,
    site_data_path: Path = SITE_DATA_PATH,
    deck_content_path: Path = DECK_CONTENT_PATH,
    prose_path: Path = PROSE_PATH,
    run_site_build: bool = True,
) -> SubmissionResult:
    root = root.resolve()
    output_dir = output_dir.resolve()
    phase4_path = phase4_path.resolve()
    figures_dir = figures_dir.resolve()
    site_content_path = site_content_path.resolve()
    site_data_path = site_data_path.resolve()
    deck_content_path = deck_content_path.resolve()
    prose_path = prose_path.resolve()

    phase4 = _load_phase4(phase4_path)
    site_content = _load_json(site_content_path)
    site_data = _load_json(site_data_path)
    if not deck_content_path.exists():
        raise FileNotFoundError(f"Deck content missing: {deck_content_path}")
    figures = sorted(figures_dir.glob("*.png"))
    if not figures:
        raise FileNotFoundError(f"No figure PNGs found in {figures_dir}")

    output_dir.mkdir(parents=True, exist_ok=True)
    report_path = output_dir / "ESG_Momentum_Engine_Report.docx"
    deck_path = output_dir / "ESG_Momentum_Engine.pptx"
    site_path = output_dir / "site"

    report_path, remaining_markers = _build_report(root, report_path, phase4_path, prose_path)
    generated_site_path = _run_site_build(root / "site", site_path) if run_site_build else None
    _build_deck(deck_content_path, deck_path)

    result = SubmissionResult(
        output_dir=output_dir,
        report_path=report_path,
        deck_path=deck_path,
        site_path=generated_site_path,
        readme_path=output_dir / "README.md",
        remaining_markers=remaining_markers,
        figure_count=len(figures),
        site_build_ran=run_site_build,
    )
    readme_path = _write_package_readme(output_dir, result, site_content, site_data, phase4)
    result.readme_path = readme_path
    return result


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Assemble final ESG Momentum Engine submission package.")
    parser.add_argument("--root", type=Path, default=PROJECT_ROOT, help="Repository root")
    parser.add_argument("--output-dir", type=Path, default=SUBMISSION_DIR, help="Submission output directory")
    parser.add_argument("--phase4", type=Path, default=PHASE4_PATH, help="Path to phase4_results.pkl")
    parser.add_argument("--figures-dir", type=Path, default=FIGURES_DIR, help="Path to figure PNG directory")
    parser.add_argument("--site-content", type=Path, default=SITE_CONTENT_PATH, help="Path to site_content.json")
    parser.add_argument("--site-data", type=Path, default=SITE_DATA_PATH, help="Path to site_data.json")
    parser.add_argument("--deck-content", type=Path, default=DECK_CONTENT_PATH, help="Path to deck_content.yaml")
    parser.add_argument("--prose", type=Path, default=PROSE_PATH, help="Path to submission_prose.yaml")
    parser.add_argument("--skip-site-build", action="store_true", help="Skip npm build for smoke tests")
    args = parser.parse_args(argv)

    try:
        result = assemble_submission(
            root=args.root,
            output_dir=args.output_dir,
            phase4_path=args.phase4,
            figures_dir=args.figures_dir,
            site_content_path=args.site_content,
            site_data_path=args.site_data,
            deck_content_path=args.deck_content,
            prose_path=args.prose,
            run_site_build=not args.skip_site_build,
        )
    except Exception as exc:
        print(f"FAIL: {exc}", file=sys.stderr)
        return 1

    print(f"PASS: submission package written to {result.output_dir}")
    print(f"report={result.report_path}")
    print(f"deck={result.deck_path}")
    print(f"site={result.site_path if result.site_path else 'skipped'}")
    print(f"readme={result.readme_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
